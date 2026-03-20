"""DOCX 解析器：文本 + 图片提取。"""

from __future__ import annotations

from pathlib import Path

import structlog

from app.agents.ingest.parse_utils import save_image_bytes
from app.core.exceptions import FileParseError

logger = structlog.get_logger()


async def parse_docx(file_path: str | Path, asset_dir: Path) -> str:
    """DOCX 解析主入口：MarkItDown 优先，python-docx 兜底。"""
    path = Path(file_path)
    logger.info("parse_docx_start", filename=path.name)

    # 先尝试 MarkItDown
    try:
        from markitdown import MarkItDown

        result = MarkItDown().convert(str(path))
        if result.text_content and result.text_content.strip():
            _extract_docx_images(path, asset_dir)
            return result.text_content
    except Exception as exc:
        logger.warning("parse_docx_markitdown_failed", filename=path.name, error=str(exc))

    # Fallback: python-docx 手动提取
    try:
        return _parse_docx_native(path, asset_dir)
    except Exception as exc:
        logger.warning("parse_docx_native_failed", filename=path.name, error=str(exc))

    raise FileParseError(path.name, reason="DOCX 解析失败。")


# ---------------------------------------------------------------------------
# python-docx 原生解析
# ---------------------------------------------------------------------------


def _parse_docx_native(file_path: Path, asset_dir: Path) -> str:
    """用 python-docx 手动提取文本和图片。"""
    from docx import Document

    doc = Document(str(file_path))
    sections: list[str] = []
    image_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检测标题样式
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading", "").strip())
            except ValueError:
                level = 2
            sections.append(f"{'#' * level} {text}")
        else:
            sections.append(text)

        # 检查段落中的内嵌图片
        has_drawing = False
        for run in para.runs:
            drawings = run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            if drawings:
                has_drawing = True
                break

        if has_drawing:
            for rel in para.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        img_bytes = rel.target_part.blob
                        if len(img_bytes) < 1024:
                            continue
                        ext = Path(rel.target_ref).suffix or ".png"
                        image_count += 1
                        filename = save_image_bytes(
                            img_bytes, asset_dir,
                            name_hint=f"doc_img{image_count}",
                            ext=ext,
                        )
                        sections.append(f"\n![文档图片 {image_count}]({filename})\n")
                    except Exception:
                        pass

    result = "\n".join(sections).strip()
    if not result:
        raise FileParseError(file_path.name, reason="DOCX 文本提取为空。")

    logger.info("parse_docx_native_done", filename=file_path.name, images=image_count)
    return result


# ---------------------------------------------------------------------------
# 图片补充提取
# ---------------------------------------------------------------------------


def _extract_docx_images(file_path: Path, asset_dir: Path) -> None:
    """从 DOCX 中补充提取图片（当 MarkItDown 成功时调用）。"""
    try:
        from docx import Document
    except ImportError:
        return

    doc = Document(str(file_path))
    count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                if len(img_bytes) < 1024:
                    continue
                ext = Path(rel.target_ref).suffix or ".png"
                save_image_bytes(
                    img_bytes, asset_dir,
                    name_hint=f"doc_img{count}",
                    ext=ext,
                )
                count += 1
            except Exception:
                pass
    if count > 0:
        logger.info("docx_images_supplemented", filename=file_path.name, count=count)
