"""DOCX parsers used by the ingest workflow."""

from __future__ import annotations

import asyncio
from concurrent.futures import ThreadPoolExecutor
from importlib import import_module
from importlib.util import find_spec
from pathlib import Path

import structlog

from app.shared.infra.exceptions import FileParseError
from app.workflows.ingest.parsing.docx_archive import load_docx_archive
from app.workflows.ingest.parsing.types import ParserRunOptions
from app.workflows.ingest.parsing.utils import save_image_bytes


try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    qn = None

logger = structlog.get_logger()

DOCX_NATIVE_AVAILABLE = True
DOCX_MARKITDOWN_AVAILABLE = find_spec("markitdown") is not None


async def parse_docx_with_markitdown(
    file_path: str | Path,
    asset_dir: Path,
    options: ParserRunOptions,
) -> str:
    """Parse DOCX through MarkItDown and supplement extracted assets."""

    return await asyncio.to_thread(_parse_docx_with_markitdown_sync, Path(file_path), asset_dir, options)


async def parse_docx_with_native(
    file_path: str | Path,
    asset_dir: Path,
    options: ParserRunOptions,
) -> str:
    """Parse DOCX through the best native path available in this environment."""

    return await asyncio.to_thread(_parse_docx_with_native_sync, Path(file_path), asset_dir, options)


def _parse_docx_with_markitdown_sync(path: Path, asset_dir: Path, options: ParserRunOptions) -> str:
    try:
        markitdown_module = import_module("markitdown")
        markitdown_converter = getattr(markitdown_module, "MarkItDown")
    except (AttributeError, ImportError) as exc:
        raise FileParseError(path.name, reason="MarkItDown is not available.") from exc

    logger.info("parse_docx_markitdown_start", filename=path.name, parser_parallelism=options.parser_parallelism)
    result = markitdown_converter().convert(str(path))
    if not result.text_content or not result.text_content.strip():
        raise FileParseError(path.name, reason="MarkItDown returned empty markdown.")

    if not options.skip_image_supplement:
        supplement_docx_images(
            path,
            asset_dir,
            max_images=options.asset_image_limit,
            workers=options.parser_parallelism,
            asset_name_prefix=options.asset_name_prefix,
        )
    return result.text_content


def _parse_docx_with_native_sync(path: Path, asset_dir: Path, options: ParserRunOptions) -> str:
    if Document is not None:
        try:
            return _parse_docx_with_python_docx_sync(path, asset_dir, options)
        except Exception as exc:
            logger.warning(
                "parse_docx_python_native_failed_fallback",
                filename=path.name,
                error=str(exc),
            )
    return _parse_docx_with_archive_sync(path, asset_dir, options)


def _parse_docx_with_python_docx_sync(path: Path, asset_dir: Path, options: ParserRunOptions) -> str:
    if Document is None:
        raise FileParseError(path.name, reason="python-docx is not available.")

    logger.info("parse_docx_python_native_start", filename=path.name, parser_parallelism=options.parser_parallelism)
    document = Document(str(path))
    sections: list[str] = []
    image_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(_render_paragraph(paragraph.style.name if paragraph.style else None, text))

        for image_bytes, image_ext in _extract_run_images(paragraph):
            if image_count >= options.asset_image_limit:
                logger.info(
                    "parse_docx_native_image_limit_reached",
                    filename=path.name,
                    limit=options.asset_image_limit,
                )
                break
            filename = _save_docx_image(
                image_bytes=image_bytes,
                image_ext=image_ext,
                asset_dir=asset_dir,
                image_count=image_count + 1,
                asset_name_prefix=options.asset_name_prefix,
            )
            if not filename:
                continue
            image_count += 1
            sections.append(f"![Document image {image_count}]({filename})")

    return _finalize_native_docx_result(path=path, sections=sections, image_count=image_count)


def _parse_docx_with_archive_sync(path: Path, asset_dir: Path, options: ParserRunOptions) -> str:
    logger.info("parse_docx_archive_native_start", filename=path.name, parser_parallelism=options.parser_parallelism)
    document = load_docx_archive(path)
    images_by_rel_id = {item.rel_id: item for item in document.images}
    rendered_image_rel_ids: set[str] = set()
    sections: list[str] = []
    image_count = 0

    for paragraph in document.paragraphs:
        if paragraph.text:
            sections.append(_render_paragraph(paragraph.style_name, paragraph.text))

        for rel_id in paragraph.image_rel_ids:
            if rel_id in rendered_image_rel_ids:
                continue
            if image_count >= options.asset_image_limit:
                logger.info(
                    "parse_docx_archive_image_limit_reached",
                    filename=path.name,
                    limit=options.asset_image_limit,
                )
                break
            image = images_by_rel_id.get(rel_id)
            if image is None:
                continue
            filename = _save_docx_image(
                image_bytes=image.blob,
                image_ext=Path(image.internal_path).suffix or ".png",
                asset_dir=asset_dir,
                image_count=image_count + 1,
                asset_name_prefix=options.asset_name_prefix,
            )
            if not filename:
                continue
            image_count += 1
            rendered_image_rel_ids.add(rel_id)
            sections.append(f"![Document image {image_count}]({filename})")

    return _finalize_native_docx_result(path=path, sections=sections, image_count=image_count)


def supplement_docx_images(
    file_path: Path,
    asset_dir: Path,
    *,
    max_images: int,
    workers: int,
    asset_name_prefix: str,
) -> None:
    """Extract document images even when markdown came from MarkItDown."""

    if Document is None:
        _supplement_docx_images_from_archive(
            file_path,
            asset_dir,
            max_images=max_images,
            workers=workers,
            asset_name_prefix=asset_name_prefix,
        )
        return

    document = Document(str(file_path))
    extracted_images: list[tuple[bytes, str, str]] = []
    seen_refs: set[str] = set()
    for relation in document.part.rels.values():
        if len(extracted_images) >= max_images:
            logger.info("docx_images_supplement_limited", filename=file_path.name, limit=max_images)
            break
        if "image" not in relation.reltype:
            continue
        target_ref = str(relation.target_ref)
        if target_ref in seen_refs:
            continue
        seen_refs.add(target_ref)
        image_bytes = relation.target_part.blob
        if len(image_bytes) < 1024:
            continue
        ext = Path(target_ref).suffix or ".png"
        extracted_images.append((image_bytes, ext, f"doc_img{len(extracted_images) + 1}"))

    saved_count = _save_docx_supplement_images(
        extracted_images,
        asset_dir,
        workers=workers,
        asset_name_prefix=asset_name_prefix,
    )
    if saved_count:
        logger.info("docx_images_supplemented", filename=file_path.name, count=saved_count)


def _supplement_docx_images_from_archive(
    file_path: Path,
    asset_dir: Path,
    *,
    max_images: int,
    workers: int,
    asset_name_prefix: str,
) -> None:
    document = load_docx_archive(file_path)
    extracted_images: list[tuple[bytes, str, str]] = []
    for image in document.images:
        if len(extracted_images) >= max_images:
            logger.info("docx_images_supplement_limited", filename=file_path.name, limit=max_images)
            break
        if len(image.blob) < 1024:
            continue
        extracted_images.append(
            (
                image.blob,
                Path(image.internal_path).suffix or ".png",
                f"doc_img{len(extracted_images) + 1}",
            )
        )

    saved_count = _save_docx_supplement_images(
        extracted_images,
        asset_dir,
        workers=workers,
        asset_name_prefix=asset_name_prefix,
    )
    if saved_count:
        logger.info("docx_images_supplemented", filename=file_path.name, count=saved_count)


def _save_docx_image(
    *,
    image_bytes: bytes,
    image_ext: str,
    asset_dir: Path,
    image_count: int,
    asset_name_prefix: str,
) -> str | None:
    if len(image_bytes) < 1024:
        return None
    return save_image_bytes(
        image_bytes,
        asset_dir,
        name_hint=f"doc_img{image_count}",
        ext=image_ext,
        name_prefix=asset_name_prefix,
    )


def _save_docx_supplement_images(
    images: list[tuple[bytes, str, str]],
    asset_dir: Path,
    *,
    workers: int,
    asset_name_prefix: str,
) -> int:
    if not images:
        return 0

    max_workers = min(max(workers, 1), 10)
    if max_workers == 1:
        for image_bytes, image_ext, hint in images:
            save_image_bytes(
                image_bytes,
                asset_dir,
                name_hint=hint,
                ext=image_ext,
                name_prefix=asset_name_prefix,
            )
        return len(images)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [
            executor.submit(save_image_bytes, image_bytes, asset_dir, hint, image_ext, asset_name_prefix)
            for image_bytes, image_ext, hint in images
        ]
        for future in futures:
            future.result()
    return len(images)


def _finalize_native_docx_result(
    *,
    path: Path,
    sections: list[str],
    image_count: int,
) -> str:
    result = "\n".join(part for part in sections if part).strip()
    if not result:
        raise FileParseError(path.name, reason="DOCX text extraction returned empty markdown.")

    logger.info("parse_docx_native_done", filename=path.name, images=image_count)
    return result


def _render_paragraph(style_name: str | None, text: str) -> str:
    if style_name and style_name.startswith("Heading"):
        level = _parse_heading_level(style_name)
        return f"{'#' * level} {text}"
    return text


def _parse_heading_level(style_name: str) -> int:
    suffix = style_name.replace("Heading", "", 1).strip()
    if suffix.isdigit():
        return min(max(int(suffix), 1), 6)
    return 2


def _extract_run_images(paragraph: object) -> list[tuple[bytes, str]]:
    if qn is None:
        return []

    images: list[tuple[bytes, str]] = []
    seen_ids: set[str] = set()
    for run in paragraph.runs:
        for blip in run._element.xpath(".//a:blip"):
            embed_id = blip.get(qn("r:embed"))
            if not embed_id or embed_id in seen_ids:
                continue
            seen_ids.add(embed_id)
            image_part = run.part.related_parts[embed_id]
            ext = Path(str(image_part.partname)).suffix or ".png"
            images.append((image_part.blob, ext))
    return images
