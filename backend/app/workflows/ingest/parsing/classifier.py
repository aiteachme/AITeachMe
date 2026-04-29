"""Lightweight ingest classification used to plan parser routing."""

from __future__ import annotations

import re
from pathlib import Path

import structlog
from pydantic import BaseModel, Field

from app.workflows.ingest.parsing.docx_archive import summarize_docx_archive
from app.workflows.ingest.parsing.formats import (
    categorize_text_extension,
    is_text_extension,
    normalize_extension,
)
from app.workflows.ingest.parsing.text import is_probably_text_file, read_text_file

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None


logger = structlog.get_logger()

_ZH_RE = re.compile(r"[\u4e00-\u9fff]")
_EN_RE = re.compile(r"[a-zA-Z]")
_HEADING_LIKE_RE = re.compile(
    r"^(?:\u7b2c[\u4e00-\u9fff0-9]+[\u7ae0\u8282\u7bc7]|Chapter\s+\d+|Section\s+\d+|\d+[\.\s])",
    re.MULTILINE,
)
_MARKDOWN_HEADING_RE = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_TABLE_LIKE_RE = re.compile(r"\|.*\|.*\||\+[-=]+\+", re.MULTILINE)
_FORMULA_RE = re.compile(r"(?:\\frac|\\sum|\\int|\\sqrt|\u2211|\u221a|\u2264|\u2265)")


class ClassificationResult(BaseModel):
    """Classification output consumed by the ingest workflow."""

    file_category: str
    text_density: float = 0.0
    ocr_ratio: float = 0.0
    image_page_ratio: float = 0.0
    heading_count: int = 0
    estimated_pages: int = 0
    detected_language: str = "unknown"
    has_tables: bool = False
    has_formulas: bool = False
    recommended_parser: str = ""
    fallback_parsers: list[str] = Field(default_factory=list)

    def to_dict(self) -> dict[str, object]:
        return self.model_dump()


def classify_file(file_path: str | Path, filetype: str) -> ClassificationResult:
    """Inspect a file quickly and choose a parser chain."""

    path = Path(file_path)
    extension = normalize_extension(filetype)

    if extension == ".pdf":
        return _classify_pdf(path)
    if extension == ".docx":
        return _classify_docx(path)
    if is_text_extension(extension):
        return _classify_text_file(path, extension)
    if is_probably_text_file(path):
        return _classify_text_file(path, extension)
    return ClassificationResult(
        file_category="unknown",
        recommended_parser="markitdown",
    )


def _classify_pdf(path: Path) -> ClassificationResult:
    result = ClassificationResult(
        file_category="pdf",
        recommended_parser="markitdown",
        fallback_parsers=[],
    )
    logger.info(
        "classify_pdf_done",
        filename=path.name,
        category=result.file_category,
        parser=result.recommended_parser,
    )
    return result


def _classify_docx(path: Path) -> ClassificationResult:
    try:
        paragraph_count, heading_count, avg_density, sample_text = _probe_docx_text(path)
    except Exception:
        paragraph_count = 0
        heading_count = 0
        avg_density = 0.0
        sample_text = ""

    return ClassificationResult(
        file_category="docx",
        text_density=round(avg_density, 1),
        heading_count=heading_count,
        estimated_pages=max(paragraph_count // 30, 1) if paragraph_count else 0,
        detected_language=_detect_language(sample_text[:5000]) if sample_text else "unknown",
        has_tables=bool(_TABLE_LIKE_RE.search(sample_text[:10000])),
        has_formulas=bool(_FORMULA_RE.search(sample_text[:10000])),
        recommended_parser="markitdown",
        fallback_parsers=["docx_native"],
    )


def _probe_docx_text(path: Path) -> tuple[int, int, float, str]:
    if Document is not None:
        try:
            document = Document(str(path))
            texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            heading_count = sum(
                1
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name and paragraph.style.name.lower().startswith("heading")
            )
            paragraph_count = len(texts)
            total_text = sum(len(item) for item in texts)
            return (
                paragraph_count,
                heading_count,
                total_text / max(paragraph_count, 1),
                "\n".join(texts[:80])[:10000],
            )
        except Exception:
            pass

    summary = summarize_docx_archive(path)
    return (
        summary.paragraph_count,
        summary.heading_count,
        summary.total_text_chars / max(summary.paragraph_count, 1) if summary.paragraph_count else 0.0,
        summary.sample_text,
    )


def _classify_text_file(path: Path, extension: str) -> ClassificationResult:
    try:
        text = read_text_file(path)
    except Exception:
        text = ""

    stripped = text.strip()
    non_empty_lines = [line for line in text.splitlines() if line.strip()]
    line_count = len(non_empty_lines)
    heading_count = len(_HEADING_LIKE_RE.findall(text[:10000])) + len(_MARKDOWN_HEADING_RE.findall(text[:10000]))
    estimated_pages = max(line_count // 50, 1) if line_count else 0
    file_category = categorize_text_extension(extension)
    return ClassificationResult(
        file_category=file_category,
        text_density=round(len(stripped) / max(line_count, 1), 1) if stripped else 0.0,
        heading_count=heading_count,
        estimated_pages=estimated_pages,
        detected_language=_detect_language(text[:5000]) if stripped else "unknown",
        has_tables=bool(_TABLE_LIKE_RE.search(text[:10000])),
        has_formulas=bool(_FORMULA_RE.search(text[:10000])),
        recommended_parser="text_native",
    )


def _detect_language(text: str) -> str:
    zh_count = len(_ZH_RE.findall(text))
    en_count = len(_EN_RE.findall(text))
    if zh_count > en_count * 2:
        return "zh"
    if en_count > zh_count * 2:
        return "en"
    return "mixed"
