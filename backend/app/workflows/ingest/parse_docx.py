"""DOCX parsers used by the ingest workflow."""

from __future__ import annotations

from pathlib import Path
import structlog

from app.core.exceptions import FileParseError
from app.workflows.ingest.parse_utils import save_image_bytes


try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    qn = None

try:
    from markitdown import MarkItDown
except ImportError:
    MarkItDown = None


logger = structlog.get_logger()


async def parse_docx_with_markitdown(file_path: str | Path, asset_dir: Path) -> str:
    """Parse DOCX through MarkItDown and supplement extracted assets."""

    path = Path(file_path)
    if MarkItDown is None:
        raise FileParseError(path.name, reason="MarkItDown is not available.")

    logger.info("parse_docx_markitdown_start", filename=path.name)
    result = MarkItDown().convert(str(path))
    if not result.text_content or not result.text_content.strip():
        raise FileParseError(path.name, reason="MarkItDown returned empty markdown.")

    supplement_docx_images(path, asset_dir)
    return result.text_content


async def parse_docx_with_python_docx(file_path: str | Path, asset_dir: Path) -> str:
    """Parse DOCX through python-docx with image placeholders."""

    path = Path(file_path)
    if Document is None:
        raise FileParseError(path.name, reason="python-docx is not available.")

    logger.info("parse_docx_python_native_start", filename=path.name)
    document = Document(str(path))
    sections: list[str] = []
    image_count = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(_render_paragraph(paragraph.style.name if paragraph.style else None, text))

        for image_bytes, image_ext in _extract_run_images(paragraph):
            if len(image_bytes) < 1024:
                continue
            image_count += 1
            filename = save_image_bytes(
                image_bytes,
                asset_dir,
                name_hint=f"doc_img{image_count}",
                ext=image_ext,
            )
            sections.append(f"![Document image {image_count}]({filename})")

    result = "\n".join(part for part in sections if part).strip()
    if not result:
        raise FileParseError(path.name, reason="DOCX text extraction returned empty markdown.")

    logger.info("parse_docx_python_native_done", filename=path.name, images=image_count)
    return result


def supplement_docx_images(file_path: Path, asset_dir: Path) -> None:
    """Extract document images even when markdown came from MarkItDown."""

    if Document is None:
        return

    document = Document(str(file_path))
    count = 0
    seen_refs: set[str] = set()
    for relation in document.part.rels.values():
        if "image" not in relation.reltype:
            continue
        target_ref = str(relation.target_ref)
        if target_ref in seen_refs:
            continue
        seen_refs.add(target_ref)
        image_bytes = relation.target_part.blob
        if len(image_bytes) < 1024:
            continue
        ext = Path(target_ref).suffix or ".png"
        save_image_bytes(
            image_bytes,
            asset_dir,
            name_hint=f"doc_img{count + 1}",
            ext=ext,
        )
        count += 1

    if count:
        logger.info("docx_images_supplemented", filename=file_path.name, count=count)


def _render_paragraph(style_name: str | None, text: str) -> str:
    if style_name and style_name.startswith("Heading"):
        level = _parse_heading_level(style_name)
        return f"{'#' * level} {text}"
    return text


def _parse_heading_level(style_name: str) -> int:
    suffix = style_name.replace("Heading", "", 1).strip()
    if suffix.isdigit():
        return min(max(int(suffix), 1), 6)
    return 2


def _extract_run_images(paragraph: object) -> list[tuple[bytes, str]]:
    if qn is None:
        return []

    images: list[tuple[bytes, str]] = []
    seen_ids: set[str] = set()
    for run in paragraph.runs:
        for blip in run._element.xpath(".//a:blip"):
            embed_id = blip.get(qn("r:embed"))
            if not embed_id or embed_id in seen_ids:
                continue
            seen_ids.add(embed_id)
            image_part = run.part.related_parts[embed_id]
            ext = Path(str(image_part.partname)).suffix or ".png"
            images.append((image_part.blob, ext))
    return images
